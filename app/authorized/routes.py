# -*- encoding: utf-8 -*-
"""

"""

from cmath import log
from email import message
from importlib.machinery import all_suffixes
from operator import gt
from flask import jsonify, render_template, redirect, request, url_for, send_file, session
from jinja2 import TemplateNotFound
# from flask_login import (
#     current_user,
#     login_required,
#     login_user,
#     logout_user
# )
from docx import Document
import json
from app import db
from app.authorized import authorized
from app.authorized.forms import bpForm, fmForm, ConsultingForm, MaintenanceForm
import logging
import os
from app.authorized.module_locator import module_path
from app.base.models import User, login_required
from datetime import date, datetime, timedelta

def path_locator():
   pass 

cwd = module_path(path_locator)
cwd = cwd.replace("/routes.py", "")


# from app.base.models import User

# from app.base.util import verify_pass

# STATISTICS ROUTES #
@authorized.route('/visitors_info')
@login_required
def get_all_visits():
    try:
        all_information = db.statistics.find()
        now_time = datetime.now() # or now??
        today_info_range = now_time + timedelta(days=-1)
        week_info_range =  now_time + timedelta(days=-7)
        month_info_range = now_time +timedelta(days=-31)
        
        filtered_today = [p for p in all_information if p["date_obj"] > today_info_range]
        filtered_month = [p for p in db.statistics.find() if p["date_obj"] > month_info_range]
        filtered_week = [p for p in db.statistics.find() if p["date_obj"] > week_info_range]
        
        return render_template("clients_info/statistics.html",
                                # db.statistics.find() -->filtered_month
                                 all_info = sorted(db.statistics.find(), key = lambda x: x["date_obj"], reverse=True),
                                 today_info=filtered_today,
                                 week_info=filtered_week,
                                 month_info=filtered_month,
                                 description="all"
                                 
                                 )
    except Exception as e :
        # print(e)
        return "Error on the server side"

@authorized.route('/visitors_info/<page_type>')
@login_required
def get_page_visits(page_type):
    try:
        all_information = db.statistics.find({"page": page_type})
        now_time = datetime.now() # or now??
        today_info_range = now_time + timedelta(days=-1)
        week_info_range =  now_time + timedelta(days=-7)
        month_info_range = now_time +timedelta(days=-31)
        
        filtered_today = [p for p in all_information if p["date_obj"] > today_info_range]
        filtered_month = [p for p in db.statistics.find({"page": page_type}) if p["date_obj"] > month_info_range]
        filtered_week = [p for p in db.statistics.find({"page": page_type}) if p["date_obj"] > week_info_range]
        print(filtered_today)
        # print(filtered_today)
        return render_template("clients_info/statistics.html",
                                #  instead of db.statistics.find({...}) -- filtered_month
                                all_info = sorted(db.statistics.find({"page": page_type}), key = lambda x: x["date_obj"], reverse=True),
                                 today_info=filtered_today,
                                 week_info=filtered_week,
                                 month_info=filtered_month,
                                 description=page_type
                                 
                                 )
        all_information = db.statistics.find({"page": page_type})
        return render_template("clients_info/statistics.html",
                                 all_info = all_information, 
                                 description=page_type )
    except Exception as e:
        print(e)
        return "Error on the server side"
    


####### CHANGE USERS STATUS PAGE ######
@authorized.route("/35ea0473", methods=['POST', 'GET'])
@login_required
def users_status():
    if request.method=='POST':
        
        to_change_user = request.form['username']
        to_change_status = request.form['new-status']
        try:
            db.users.update_one({"username": to_change_user}, {"$set": {"payment_status": to_change_status}})
            users = list(db.users.find())
            return render_template("/clients_info/all_users.html", users=users)
        except Exception as e:
            # print(e)
            return "Произошла ошибка, попробуйте ввести еще раз"
        return "CHECK Print"
    else:
        try:
            users = list(db.users.find())
            all_users = []
            return render_template("/clients_info/all_users.html", users=users)
        except Exception as e:
            return "Произошла ошибка на стороне сервера и данные не были получены"

################################################


@authorized.route('/index')
@login_required
def route_index():
    now = datetime.now()
    # dd/mm/YY H:M:S
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    username = session['user']
    click_info = {
        "username":username,
        "access_time":dt_string,
        "page": "index",
        "date_obj": now
    }
    db.insert(db.statistics, click_info)
    return render_template("index.html")

@authorized.route('/<number>_lesson', methods=['POST', 'GET'])
@login_required
def route_lesson(number):
    username = session['user']
    mes1 = "У вас открылась главная страница, потому что пока нет доступа к той странице, на которую вы хотели зайти."
    mes2 = "Если у вас нет доступа к странице, но должен быть - свяжитесь с менеджером."
    user = db.get_one(db.users, "username", username)


    user_status = user['payment_status']
    # код отвечающий за обратную связь и перевод на следующий урок:
    if request.method=='POST':
        user_feedbacks = db.feedbacks.find_one({"username":username})
        if user_feedbacks:
            if number not in user_feedbacks["lessons_feedbacks"].keys():
                print("\n\nUSER DID NOT left the comments on this lesson!\n\n")
                rating = None
                if request.form.get("perfect"):
                    rating = "perfect"
                if request.form.get("good"):
                    rating = "good"
                if request.form.get("normal"):
                    rating = "normal"
                if request.form.get("notgood"):
                    rating = "notgood"
                if request.form.get("bad"):
                    rating = "bad"
                comment = request.form["modal-comment"]
                feedbacks_dictionary = user_feedbacks["lessons_feedbacks"] 
                # adding new lesson value to the dictionary
                feedbacks_dictionary[number] = {
                    "rating": rating,
                    "comments":comment
                }
                # db.users.update_one({"username": username}, {"$set": {"payment_status": to_change_status}})
                db.feedbacks.update_one({"username": username}, {"$set": {"lessons_feedbacks": feedbacks_dictionary}})
                # db.feedbacks.find_and_modify(query={"username":username}, update={"lessons_feedbacks":feedbacks_dictionary})
        else:
            rating = None
            if request.form.get("perfect"):
                rating = "perfect"
            if request.form.get("good"):
                   rating = "good"
            if request.form.get("normal"):
                rating = "normal"
            if request.form.get("notgood"):
                rating = "notgood"
            if request.form.get("bad"):
                rating = "bad"

            comment = request.form["modal-comment"]
            print("CREATING NEW FEEDBACK!!\n\n")
            feedback_info = {
                "username": username,
                "lessons_feedbacks" : {
                    number : {
                        "rating" : rating,
                        "comments": comment
                    }
                    
                }
            }
            db.insert(db.feedbacks, feedback_info)
            # print("\nELSE: \n" )
            # print(db.get_one(db.feedbacks, "username", username))
        # code for checking if correct
        # print(db.get_one(db.feedbacks, "username", username))
        if number == "zero":
            number = "first"
        elif number=="first":
            number = "second"
        elif number=="second":
            number = "third"
        elif number=="third":
            number = "fourth"
        elif number=="fourth":
            number = "fifth"
        elif number == "fifth":
            number = "sixth"
        elif number=="sixth":
            number = "sevens"
        elif number=="sevens":
            number = "eights"
        elif number=="eights":
            number="bonus"
        elif number=="bonus":
            number = "bonus2"
        
        return redirect("/authorized/"+number+"_lesson")
        ##### END OF THE CODE ###

    if request.method=='GET':
    #statistics code
        if user_status == "course-pro" and user_status == "course-standard":
            now = datetime.now()
            # dd/mm/YY H:M:S
            dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
            click_info = {
                "username":username,
                "access_time":dt_string,
                "page": (number + "_lesson"),
                "date_obj": now
            }
            db.insert(db.statistics, click_info)
        ###############################

    
        user_fdb_hst = db.get_one(db.feedbacks, "username", username)
        logging.error(f"AHTUNG AJTUNG {user_fdb_hst}")

        if not user_fdb_hst:
            feedback_info = {
                "username": username,
                "lessons_feedbacks" : {}
            }
            db.insert(db.feedbacks, feedback_info)
            
        user_fdb_hst = db.get_one(db.feedbacks, "username", username)
        logging.error(len(user_fdb_hst["lessons_feedbacks"]))
        
        if user_status == "super-admin":
            return render_template(f"dkb-lessons/{number}_lesson.html")
        if user_status != "course-pro" and user_status != "course-standard":
            return render_template(f"index.html", message1 = mes1, message2 = mes2)
        if number == "bonus2" and user_status == "course-standard":
            return render_template(f"index.html", message1 = mes1, message2 = mes2)
        
        checked_rating = False
        print("\n\n")
        print(user_fdb_hst)
        # print("\n\n")
        mes1 = "Вы должны посмотреть прошлые уроки и оставить свою оценку на уроках, чтобы открылся доступ к следующим урокам"
        if number == "zero" and (user_status == "course-pro" or user_status == "course-standard"):
            
            if user_fdb_hst and (len(user_fdb_hst["lessons_feedbacks"]) >= 1):
                checked_rating = True
            print(checked_rating)
            return render_template(f"dkb-lessons/zero_lesson.html",checked_rating=checked_rating)
             
        if number == "first" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=1:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 1):
                checked_rating = True
            return render_template(f"dkb-lessons/first_lesson.html", checked_rating=checked_rating)
        if number == "second" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=2:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 2):
                checked_rating = True
            return render_template(f"dkb-lessons/second_lesson.html",  checked_rating=checked_rating)
        if number == "third" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=3:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 3):
                checked_rating = True
            return render_template(f"dkb-lessons/third_lesson.html", checked_rating=checked_rating)
        if number == "fourth" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=4:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 4):
                checked_rating = True
            return render_template(f"dkb-lessons/fourth_lesson.html", checked_rating=checked_rating)
        if number == "fifth" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=5:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 5):
                checked_rating = True
            return render_template(f"dkb-lessons/fifth_lesson.html", checked_rating=checked_rating)
        if number == "sixth" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=6:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 6):
                checked_rating = True
            return render_template(f"dkb-lessons/sixth_lesson.html", checked_rating=checked_rating)
        if number == "sevens" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=7:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 7):
                checked_rating = True
            return render_template(f"dkb-lessons/sevens_lesson.html", checked_rating=checked_rating)
        if number == "eights" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=8:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 8):
                checked_rating = True
            return render_template(f"dkb-lessons/eights_lesson.html", checked_rating=checked_rating)
        if number == "bonus" and (user_status == "course-pro" or user_status == "course-standard") and len(user_fdb_hst["lessons_feedbacks"])>=9:
            if (len(user_fdb_hst["lessons_feedbacks"]) > 9):
                checked_rating = True
            return render_template(f"dkb-lessons/bonus_lesson.html", checked_rating=checked_rating)
        if number == "bonus2" and user_status == "course-pro" and len(user_fdb_hst["lessons_feedbacks"])>=10 :
            if (len(user_fdb_hst["lessons_feedbacks"]) > 10):
                checked_rating = True
            return render_template(f"dkb-lessons/bonus2_lesson.html")
        return render_template(f"index.html", message1 = mes1, message2 = mes2)
        # if number != "first" and (user_status == "course-pro" or user_status == "course-standard"):
        #     return render_template(f"index.html", message1 = mes1, message2 = mes2)



@authorized.route('/change_status_<username>_<new_status>')
def route_change_status(username, new_status):
    to_change_user = username
    to_change_status = new_status
    db.users.update_one({"username": username}, {"$set": {"payment_status": to_change_status}})
    return "done"

@authorized.route('/get_info_<username>')
def route_get_info(username):
    this_user = db.get_one(db.users, "username", username)
    if this_user:
        return this_user['payment_status']
    else:
        return "no such user"

@authorized.route('/delete_<username>')
def route_delete(username):
    db.delete(db.users, "username", username)
    return "done"

@authorized.route('/get_all')
def route_get_all():
    users = db.get_all(db.users)
    all_users = {}
    a = 1
    for user in users:
        all_users[a] = user
        a = a + 1
    return all_users


@authorized.route('/get_all_consulting')
def route_get_all_consulting():
    users = db.get_all(db.consulting)
    all_users = {}
    a = 1
    for user in users:
        all_users[a] = user
        a = a + 1
    return all_users


@authorized.route('/get_all_maintenance')
def route_get_all_maintenance():
    users = db.get_all(db.maintenance)
    all_users = {}
    a = 1
    for user in users:
        all_users[a] = user
        a = a + 1
    return all_users

@authorized.route('/confidentiality_politics')
def route_confidentiality_politics():
    return render_template("dkb-lessons/confidentiality_politics.html")


@authorized.route('/marafon_timer', methods=['GET'])
def route_marafon_timer():
    return "15:29:44"

##########################################
# ADD TWO PAGES WITH CLIENTS INFORMATION
@authorized.route('/clients_info_maint', methods=['GET'])
def clients_maintenance():
    try: 
        clients_maintenance = db.maintenance.find()
        return render_template("clients_info/maint.html", clients=clients_maintenance)
    except Exception as e:
        return "Ошибка на стороне сервера"
    

@authorized.route('/clients_info_consult', methods=['GET'])
def clients_consulting():
    try: 
        clients_consulting = db.consulting.find()
        return render_template("clients_info/consult.html", clients=clients_consulting)
    except Exception as e:
        return "Ошибка на стороне сервера"

##########################################

@authorized.route('/maintenance', methods=['GET', 'POST'])
def route_maintenance():

    
    maintenance_form = MaintenanceForm(request.form)
    if request.method == "GET":
        # statistics 
        now = datetime.now()
        # dd/mm/YY H:M:S
        dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
        
        click_info = {
            "username":"incognito",
            "access_time":dt_string,
            "page": "maintenance",
            "date_obj": now
        }
        db.insert(db.statistics, click_info)
        ###############
        return render_template("grant_page/maintenance.html", form = maintenance_form, message = "")
    if request.method == "POST":
        name = request.form['name']
        phone_number = request.form['phone_number']
        phone_number_again = request.form['phone_number_again']
        have_idea = request.form['have_idea']
        your_sphere = request.form['your_sphere']
        if phone_number != phone_number_again:
            return render_template("grant_page/maintenance.html", form = maintenance_form, message = "Номера не совпадают")
        else:
            person_info = {
                "name": name,
                "phone_number": phone_number,
                "have_idea": have_idea,
                "your_sphere": your_sphere
            }
            db.insert(db.maintenance, person_info)
            return render_template("grant_page/thanks.html")


@authorized.route('/consulting', methods=['GET', 'POST'])
def route_consulting():
    consulting_form = ConsultingForm(request.form)
    if request.method == "GET":
        # statistics 
        now = datetime.now()
        # dd/mm/YY H:M:S
        dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
        
        click_info = {
            "username":"incognito",
            "access_time":dt_string,
            "page": "consulting",
            "date_obj": now
        }
        db.insert(db.statistics, click_info)
        ###############
        return render_template("grant_page/consulting.html", form = consulting_form, message = "")
    if request.method == "POST":
        name = request.form['name']
        phone_number = request.form['phone_number']
        phone_number_again = request.form['phone_number_again']
        question = request.form['question']
        if phone_number != phone_number_again:
            return render_template("grant_page/consulting.html", form = consulting_form, message = "Номера не совпадают")
        else:
            person_info = {
                "name": name,
                "phone_number": phone_number,
                "question": question
            }
            db.insert(db.consulting, person_info)
            return render_template("grant_page/thanks.html")



@authorized.route('/marafon', methods=['GET'])
def route_marafon():
    if request.method == "GET":
        # statistics 
        now = datetime.now()
        # dd/mm/YY H:M:S
        dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
        
        click_info = {
            "username":"incognito",
            "access_time":dt_string,
            "page": "marafon",
            "date_obj": now
        }
        db.insert(db.statistics, click_info)
        ###############
        return render_template("grant_page/marafon.html")


@authorized.route('/payment_links', methods=['GET'])
def route_payment_links():
    if request.method == "GET":
        return render_template("grant_page/payment_links.html")


@authorized.route('/oferta')
def route_oferta():
    return render_template("dkb-lessons/oferta.html")

@authorized.route('/form')
def route_form():
    return render_template("dkb-lessons/form.html")

@authorized.route('/prices')
def route_prices():
    return render_template("dkb-lessons/prices.html")


@authorized.route('/general_data_saver', methods=['POST'])
@login_required
def general_data_saver():
    bp_form = bpForm(request.form)
    project_name = request.form["project_name"]
    businessman_name = request.form["businessman_name"]
    capital_distribution = request.form["capital_distribution"]
    address = request.form["address"]
    website = request.form["website"]
    self_investment = request.form["self_investment"]
    your_product = request.form["your_product"]
    start_deadline = request.form["start_deadline"]
    username = session['user']
    test = db.get_one(db.user_bp, "username", username)
    if not test:
        db.user_bp.insert_one({"username": username})
    db.user_bp.update_one({"username": username}, {"$set": {"project_name": project_name}})
    db.user_bp.update_one({"username": username}, {"$set": {"businessman_name": businessman_name}})
    db.user_bp.update_one({"username": username}, {"$set": {"capital_distribution": capital_distribution}})
    db.user_bp.update_one({"username": username}, {"$set": {"address": address}})
    db.user_bp.update_one({"username": username}, {"$set": {"website": website}})
    db.user_bp.update_one({"username": username}, {"$set": {"self_investment": self_investment}})
    db.user_bp.update_one({"username": username}, {"$set": {"your_product": your_product}})
    db.user_bp.update_one({"username": username}, {"$set": {"start_deadline": start_deadline}})
    return render_template('bp_constructor.html', form = bp_form)


@authorized.route('/technological_data_saver', methods=['POST'])
@login_required
def technological_data_saver():
    bp_form = bpForm(request.form)
    problem = request.form["problem"]
    solution = request.form["solution"]
    innovation = request.form["innovation"]
    patents = request.form["patents"]
    business_model = request.form["business_model"]
    resources = request.form["resources"]
    username = session['user']
    test = db.get_one(db.user_bp, "username", username)
    if not test:
        db.user_bp.insert_one({"username": username})
    db.user_bp.update_one({"username": username}, {"$set": {"problem": problem}})
    db.user_bp.update_one({"username": username}, {"$set": {"solution": solution}})
    db.user_bp.update_one({"username": username}, {"$set": {"innovation": innovation}})
    db.user_bp.update_one({"username": username}, {"$set": {"patents": patents}})
    db.user_bp.update_one({"username": username}, {"$set": {"business_model": business_model}})
    db.user_bp.update_one({"username": username}, {"$set": {"resources": resources}})
    return render_template('bp_constructor.html', form = bp_form)


@authorized.route('/commercialization_data_saver', methods=['POST'])
@login_required
def commercialization_data_saver():
    bp_form = bpForm(request.form)
    analysis = request.form["analysis"]
    competitors = request.form["competitors"]
    strategy = request.form["strategy"]

    competitor1 = request.form["competitor1"]
    competitor2 = request.form["competitor2"]
    competitor3 = request.form["competitor3"]
    competitor1value1 = request.form["competitor1value1"]
    competitor2value1 = request.form["competitor2value1"]
    competitor3value1 = request.form["competitor3value1"]
    competitor1value2 = request.form["competitor1value2"]
    competitor2value2 = request.form["competitor2value2"]
    competitor3value2 = request.form["competitor3value2"]
    competitor1value3 = request.form["competitor1value3"]
    competitor2value3 = request.form["competitor2value3"]
    competitor3value3 = request.form["competitor3value3"]
    competitor1value4 = request.form["competitor1value4"]
    competitor2value4 = request.form["competitor2value4"]
    competitor3value4 = request.form["competitor3value4"]
    advantage1 = request.form["advantage1"]
    advantage2 = request.form["advantage2"]
    advantage3 = request.form["advantage3"]
    advantage4 = request.form["advantage4"]

    username = session['user']
    test = db.get_one(db.user_bp, "username", username)
    if not test:
        db.user_bp.insert_one({"username": username})
    db.user_bp.update_one({"username": username}, {"$set": {"analysis": analysis}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitors": competitors}})
    db.user_bp.update_one({"username": username}, {"$set": {"strategy": strategy}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor1": competitor1}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor2": competitor2}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor3": competitor3}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor1value1": competitor1value1}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor2value1": competitor2value1}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor3value1": competitor3value1}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor1value2": competitor1value2}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor2value2": competitor2value2}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor3value2": competitor3value2}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor1value3": competitor1value3}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor2value3": competitor2value3}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor3value3": competitor3value3}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor1value4": competitor1value4}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor2value4": competitor2value4}})
    db.user_bp.update_one({"username": username}, {"$set": {"competitor3value4": competitor3value4}})
    db.user_bp.update_one({"username": username}, {"$set": {"advantage1": advantage1}})
    db.user_bp.update_one({"username": username}, {"$set": {"advantage2": advantage2}})
    db.user_bp.update_one({"username": username}, {"$set": {"advantage3": advantage3}})
    db.user_bp.update_one({"username": username}, {"$set": {"advantage4": advantage4}})
    return render_template('bp_constructor.html', form = bp_form)




@authorized.route('/team_data_saver', methods=['POST'])
@login_required
def team_data_saver():
    bp_form = bpForm(request.form)
    team = request.form["team"]
    experience = request.form["experience"]
    organizational_structure = request.form["organizational_structure"]
    username = session['user']
    test = db.get_one(db.user_bp, "username", username)
    if not test:
        db.user_bp.insert_one({"username": username})
    db.user_bp.update_one({"username": username}, {"$set": {"team": team}})
    db.user_bp.update_one({"username": username}, {"$set": {"experience": experience}})
    db.user_bp.update_one({"username": username}, {"$set": {"organizational_structure": organizational_structure}})
    return render_template('bp_constructor.html', form = bp_form)



@authorized.route('/project_realization_data_saver', methods=['POST'])
@login_required
def project_realization_data_saver():
    bp_form = bpForm(request.form)
    strength = request.form["strength"]
    weakness = request.form["weakness"]
    possibility = request.form["possibility"]
    threat = request.form["threat"]
    realization_plan = request.form["realization_plan"]
    username = session['user']
    test = db.get_one(db.user_bp, "username", username)
    if not test:
        db.user_bp.insert_one({"username": username})
    db.user_bp.update_one({"username": username}, {"$set": {"strength": strength}})
    db.user_bp.update_one({"username": username}, {"$set": {"weakness": weakness}})
    db.user_bp.update_one({"username": username}, {"$set": {"possibility": possibility}})
    db.user_bp.update_one({"username": username}, {"$set": {"threat": threat}})
    db.user_bp.update_one({"username": username}, {"$set": {"realization_plan": realization_plan}})
    return render_template('bp_constructor.html', form = bp_form)



@authorized.route('/financial_data_saver', methods=['POST'])
@login_required
def financial_data_saver():
    bp_form = bpForm(request.form)
    investment_before = request.form["investment_before"]
    financial_effectivity = request.form["financial_effectivity"]
    username = session['user']
    test = db.get_one(db.user_bp, "username", username)
    if not test:
        db.user_bp.insert_one({"username": username})
    db.user_bp.update_one({"username": username}, {"$set": {"investment_before": investment_before}})
    db.user_bp.update_one({"username": username}, {"$set": {"financial_effectivity": financial_effectivity}})
    return render_template('bp_constructor.html', form = bp_form)



@authorized.route('/conclusion_data_saver', methods=['POST'])
@login_required
def conclusion_data_saver():
    bp_form = bpForm(request.form)
    conclusion = request.form["conclusion"]
    username = session['user']
    test = db.get_one(db.user_bp, "username", username)
    if not test:
        db.user_bp.insert_one({"username": username})
    db.user_bp.update_one({"username": username}, {"$set": {"conclusion": conclusion}}) 
    return render_template('bp_constructor.html', form = bp_form)



@authorized.route('/bp_constructor/<action>', methods=['GET', 'POST'])
@login_required
def bp_constructor(action):
    username = session['user']
    mes1 = "У вас открылась главная страница, потому что пока нет доступа к той странице, на которую вы хотели зайти."
    mes2 = "Доступ к конструкторам будет поочередно открываться у оплативших, для более качественного обучения"
    user = db.get_one(db.users, "username", username)
    user_status = user['payment_status']
    if user_status == "beginner":
        return render_template(f"index.html", message1 = mes1, message2 = mes2)
    bp_form = bpForm(request.form)
    if request.method == 'GET' and action == "get":

        username = session['user']
        # statistics 
        now = datetime.now()
        # dd/mm/YY H:M:S
        dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
        click_info = {
            "username":username,
            "access_time":dt_string,
            "page": "bp_constructor",
            "date_obj": now
        }
        db.insert(db.statistics, click_info)
        ###############


        test = db.get_one(db.user_bp, "username", username)
        if not test:
            db.user_bp.insert_one({
            "username": username, 
            "project_name": "", 
            "businessman_name": "", 
            "capital_distribution": "", 
            "address": "",
            "website": "",
            "self_investment": "",
            "your_product": "",
            "start_deadline": "",
            "problem": "", 
            "solution": "", 
            "innovation": "", 
            "patents": "",
            "business_model": "",
            "resources": "",
            "analysis": "",
            "competitors": "",
            "strategy": "", 
            "team": "", 
            "experience": "", 
            "organizational_structure": "",
            "strength": "",
            "weakness": "",
            "possibility": "",
            "threat": "",
            "realization_plan": "",
            "investment_before": "",
            "financial_effectivity": "",
            "conclusion": "",

            "competitor1": "", 
            "competitor2": "", 
            "competitor3": "",
            "competitor1value1": "",
            "competitor2value1": "",
            "competitor3value1": "",
            "competitor1value2": "",
            "competitor2value2": "", 
            "competitor3value2": "", 
            "competitor1value3": "", 
            "competitor2value3": "",
            "competitor3value3": "",
            "competitor1value4": "",
            "competitor2value4": "",
            "competitor3value4": "",
            "advantage1": "",
            "advantage2": "",
            "advantage3": "",
            "advantage4": ""})
        bp_form.competitor1.data = db.get_one(db.user_bp, "username", username)["competitor1"]
        bp_form.competitor2.data = db.get_one(db.user_bp, "username", username)["competitor2"]
        bp_form.competitor3.data = db.get_one(db.user_bp, "username", username)["competitor3"]
        bp_form.competitor1value1.data = db.get_one(db.user_bp, "username", username)["competitor1value1"]
        bp_form.competitor2value1.data = db.get_one(db.user_bp, "username", username)["competitor2value1"]
        bp_form.competitor3value1.data = db.get_one(db.user_bp, "username", username)["competitor3value1"]
        bp_form.competitor1value2.data = db.get_one(db.user_bp, "username", username)["competitor1value2"]
        bp_form.competitor2value2.data = db.get_one(db.user_bp, "username", username)["competitor2value2"]
        bp_form.competitor3value2.data = db.get_one(db.user_bp, "username", username)["competitor3value2"]
        bp_form.competitor1value3.data = db.get_one(db.user_bp, "username", username)["competitor1value3"]
        bp_form.competitor2value3.data = db.get_one(db.user_bp, "username", username)["competitor2value3"]
        bp_form.competitor3value3.data = db.get_one(db.user_bp, "username", username)["competitor3value3"]
        bp_form.competitor1value4.data = db.get_one(db.user_bp, "username", username)["competitor1value4"]
        bp_form.competitor2value4.data = db.get_one(db.user_bp, "username", username)["competitor2value4"]
        bp_form.competitor3value4.data = db.get_one(db.user_bp, "username", username)["competitor3value4"]
        bp_form.advantage1.data = db.get_one(db.user_bp, "username", username)["advantage1"]
        bp_form.advantage2.data = db.get_one(db.user_bp, "username", username)["advantage2"]
        bp_form.advantage3.data = db.get_one(db.user_bp, "username", username)["advantage3"]
        bp_form.advantage4.data = db.get_one(db.user_bp, "username", username)["advantage4"]

        bp_form.project_name.data = db.get_one(db.user_bp, "username", username)["project_name"]
        bp_form.businessman_name.data = db.get_one(db.user_bp, "username", username)["businessman_name"]
        bp_form.capital_distribution.data = db.get_one(db.user_bp, "username", username)["capital_distribution"]
        bp_form.address.data = db.get_one(db.user_bp, "username", username)["address"]
        bp_form.website.data = db.get_one(db.user_bp, "username", username)["website"]
        bp_form.self_investment.data = db.get_one(db.user_bp, "username", username)["self_investment"]
        bp_form.your_product.data = db.get_one(db.user_bp, "username", username)["your_product"]
        bp_form.start_deadline.data = db.get_one(db.user_bp, "username", username)["start_deadline"]
        bp_form.problem.data = db.get_one(db.user_bp, "username", username)["problem"]
        bp_form.solution.data = db.get_one(db.user_bp, "username", username)["solution"]
        bp_form.innovation.data = db.get_one(db.user_bp, "username", username)["innovation"]
        bp_form.patents.data = db.get_one(db.user_bp, "username", username)["patents"]
        bp_form.business_model.data = db.get_one(db.user_bp, "username", username)["business_model"]
        bp_form.resources.data = db.get_one(db.user_bp, "username", username)["resources"]

        bp_form.analysis.data = db.get_one(db.user_bp, "username", username)["analysis"]
        bp_form.competitors.data = db.get_one(db.user_bp, "username", username)["competitors"]
        bp_form.strategy.data = db.get_one(db.user_bp, "username", username)["strategy"]

        bp_form.team.data = db.get_one(db.user_bp, "username", username)["team"]
        bp_form.experience.data = db.get_one(db.user_bp, "username", username)["experience"]
        bp_form.organizational_structure.data = db.get_one(db.user_bp, "username", username)["organizational_structure"]

        bp_form.strength.data = db.get_one(db.user_bp, "username", username)["strength"]
        bp_form.weakness.data = db.get_one(db.user_bp, "username", username)["weakness"]
        bp_form.possibility.data = db.get_one(db.user_bp, "username", username)["possibility"]
        bp_form.threat.data = db.get_one(db.user_bp, "username", username)["threat"]
        bp_form.realization_plan.data = db.get_one(db.user_bp, "username", username)["realization_plan"]

        bp_form.investment_before.data = db.get_one(db.user_bp, "username", username)["investment_before"]
        bp_form.financial_effectivity.data = db.get_one(db.user_bp, "username", username)["financial_effectivity"]

        bp_form.conclusion.data = db.get_one(db.user_bp, "username", username)["conclusion"]

        
        

        # return render_template(f"index.html", message1 = mes1, message2 = mes2)
        return render_template('bp_constructor.html', form = bp_form)
        
    elif request.method == 'GET' and action == "generate":
        project_name = db.get_one(db.user_bp, "username", username)["project_name"]
        businessman_name = db.get_one(db.user_bp, "username", username)["businessman_name"]
        capital_distribution = db.get_one(db.user_bp, "username", username)["capital_distribution"]
        address = db.get_one(db.user_bp, "username", username)["address"]
        website = db.get_one(db.user_bp, "username", username)["website"]
        self_investment = db.get_one(db.user_bp, "username", username)["self_investment"]
        your_product = db.get_one(db.user_bp, "username", username)["your_product"]
        start_deadline = db.get_one(db.user_bp, "username", username)["start_deadline"]
        problem = db.get_one(db.user_bp, "username", username)["problem"]
        solution = db.get_one(db.user_bp, "username", username)["solution"]
        innovation = db.get_one(db.user_bp, "username", username)["innovation"]
        patents = db.get_one(db.user_bp, "username", username)["patents"]
        business_model = db.get_one(db.user_bp, "username", username)["business_model"]
        resources = db.get_one(db.user_bp, "username", username)["resources"]

        analysis = db.get_one(db.user_bp, "username", username)["analysis"]
        competitors = db.get_one(db.user_bp, "username", username)["competitors"]
        strategy = db.get_one(db.user_bp, "username", username)["strategy"]

        team = db.get_one(db.user_bp, "username", username)["team"]
        experience = db.get_one(db.user_bp, "username", username)["experience"]
        organizational_structure = db.get_one(db.user_bp, "username", username)["organizational_structure"]

        strength = db.get_one(db.user_bp, "username", username)["strength"]
        weakness = db.get_one(db.user_bp, "username", username)["weakness"]
        possibility = db.get_one(db.user_bp, "username", username)["possibility"]
        threat = db.get_one(db.user_bp, "username", username)["threat"]
        realization_plan = db.get_one(db.user_bp, "username", username)["realization_plan"]

        investment_before = db.get_one(db.user_bp, "username", username)["investment_before"]
        financial_effectivity = db.get_one(db.user_bp, "username", username)["financial_effectivity"]

        conclusion = db.get_one(db.user_bp, "username", username)["conclusion"]



        competitor1 = db.get_one(db.user_bp, "username", username)["competitor1"]
        competitor2 = db.get_one(db.user_bp, "username", username)["competitor2"]
        competitor3 = db.get_one(db.user_bp, "username", username)["competitor3"]
        competitor1value1 = db.get_one(db.user_bp, "username", username)["competitor1value1"]
        competitor2value1 = db.get_one(db.user_bp, "username", username)["competitor2value1"]
        competitor3value1 = db.get_one(db.user_bp, "username", username)["competitor3value1"]
        competitor1value2 = db.get_one(db.user_bp, "username", username)["competitor1value2"]
        competitor2value2 = db.get_one(db.user_bp, "username", username)["competitor2value2"]
        competitor3value2 = db.get_one(db.user_bp, "username", username)["competitor3value2"]
        competitor1value3 = db.get_one(db.user_bp, "username", username)["competitor1value3"]
        competitor2value3 = db.get_one(db.user_bp, "username", username)["competitor2value3"]
        competitor3value3 = db.get_one(db.user_bp, "username", username)["competitor3value3"]
        competitor1value4 = db.get_one(db.user_bp, "username", username)["competitor1value4"]
        competitor2value4 = db.get_one(db.user_bp, "username", username)["competitor2value4"]
        competitor3value4 = db.get_one(db.user_bp, "username", username)["competitor3value4"]
        advantage1 = db.get_one(db.user_bp, "username", username)["advantage1"]
        advantage2 = db.get_one(db.user_bp, "username", username)["advantage2"]
        advantage3 = db.get_one(db.user_bp, "username", username)["advantage3"]
        advantage4 = db.get_one(db.user_bp, "username", username)["advantage4"]



        templ_file_path = cwd + "/templates/bp_template.docx"
        file_path = cwd + "/templates/Твой_великолепный_бизнес_план.docx"
        replace_string(templ_file_path, "PROJECT_NAME", str(project_name), file_path)
        replace_string(file_path, "PROJECT_NAMEW", str(project_name), file_path)
        replace_string(file_path, "BUSINESSMAN_NAME", str(businessman_name), file_path)
        replace_string(file_path, "CAPITAL_DISTRIBUTION", str(capital_distribution), file_path)
        replace_string(file_path, "ADDRESS", str(address), file_path)
        replace_string(file_path, "WEBSITE", str(website), file_path)
        replace_string(file_path, "SELF_INVESTMENT", str(self_investment), file_path)
        replace_string(file_path, "YOUR_PRODUCT", str(your_product), file_path)
        replace_string(file_path, "START_DEADLINE", str(start_deadline), file_path)

        replace_string(file_path, "PROBLEM", str(problem), file_path)
        replace_string(file_path, "SOLUTION", str(solution), file_path)
        replace_string(file_path, "INNOVATION", str(innovation), file_path)
        replace_string(file_path, "PATENTS", str(patents), file_path)
        replace_string(file_path, "BUSINESS_MODEL", str(business_model), file_path)
        replace_string(file_path, "RESOURCES", str(resources), file_path)

        replace_string(file_path, "ANALYSIS", str(analysis), file_path)
        replace_string(file_path, "COMPETITORS", str(competitors), file_path)


        replace_string(file_path, "competitorw1", str(competitor1), file_path)
        replace_string(file_path, "competitorw2", str(competitor2), file_path)
        replace_string(file_path, "competitorw3", str(competitor3), file_path)
        replace_string(file_path, "competitor1value1", str(competitor1value1), file_path)
        replace_string(file_path, "competitor2value1", str(competitor2value1), file_path)
        replace_string(file_path, "competitor3value1", str(competitor3value1), file_path)
        replace_string(file_path, "competitor1value2", str(competitor1value2), file_path)
        replace_string(file_path, "competitor2value2", str(competitor2value2), file_path)
        replace_string(file_path, "competitor3value2", str(competitor3value2), file_path)
        replace_string(file_path, "competitor1value3", str(competitor1value3), file_path)
        replace_string(file_path, "competitor2value3", str(competitor2value3), file_path)
        replace_string(file_path, "competitor3value3", str(competitor3value3), file_path)
        replace_string(file_path, "competitor1value4", str(competitor1value4), file_path)
        replace_string(file_path, "competitor2value4", str(competitor2value4), file_path)
        replace_string(file_path, "competitor3value4", str(competitor3value4), file_path)
        replace_string(file_path, "advantage1", str(advantage1), file_path)
        replace_string(file_path, "advantage2", str(advantage2), file_path)
        replace_string(file_path, "advantage3", str(advantage3), file_path)
        replace_string(file_path, "advantage4", str(advantage4), file_path)


        replace_string(file_path, "STRATEGY", str(strategy), file_path)

        replace_string(file_path, "TEAM", str(team), file_path)
        replace_string(file_path, "EXPERIENCE", str(experience), file_path)
        replace_string(file_path, "ORGANIZATIONAL_STRUCTURE", str(organizational_structure), file_path)

        replace_string(file_path, "STRENGTH", str(strength), file_path)
        replace_string(file_path, "WEAKNESS", str(weakness), file_path)
        replace_string(file_path, "POSSIBILITY", str(possibility), file_path)
        replace_string(file_path, "THREAT", str(threat), file_path)
        replace_string(file_path, "REALIZATION_PLAN", str(realization_plan), file_path)

        replace_string(file_path, "INVESTMENT_BEFORE", str(investment_before), file_path)
        replace_string(file_path, "FINANCIAL_EFFECTIVITY", str(financial_effectivity), file_path)
        replace_string(file_path, "CONCLUSION", str(conclusion), file_path)

        return send_file(file_path, as_attachment=True)






# @authorized.route('/add_field', methods=['POST'])
# @login_required
# def add_field():
#     pass

def is_digit(string):
    if string.isdigit():
       return True
    else:
        try:
            float(string)
            return True
        except ValueError:
            return False

@authorized.route('/initial_investments_data_saver', methods=['POST'])
@login_required
def initial_investments_data_saver():
    fm_form = fmForm(request.form)
    investment0 = request.form["investment0"]
    investment1 = request.form["investment1"]
    investment2 = request.form["investment2"]
    investment3 = request.form["investment3"]
    investment4 = request.form["investment4"]
    investment5 = request.form["investment5"]
    investment6 = request.form["investment6"]
    investment7 = request.form["investment7"]
    investment8 = request.form["investment8"]
    investment9 = request.form["investment9"]
    investment10 = request.form["investment10"]
    investment11 = request.form["investment11"]
    investment12 = request.form["investment12"]
    investment13 = request.form["investment13"]
    investment14 = request.form["investment14"]

    investment_name0 = request.form["investment_name0"]
    investment_name1 = request.form["investment_name1"]
    investment_name2 = request.form["investment_name2"]
    investment_name3 = request.form["investment_name3"]
    investment_name4 = request.form["investment_name4"]
    investment_name5 = request.form["investment_name5"]
    investment_name6 = request.form["investment_name6"]
    investment_name7 = request.form["investment_name7"]
    investment_name8 = request.form["investment_name8"]
    investment_name9 = request.form["investment_name9"]
    investment_name10 = request.form["investment_name10"]
    investment_name11 = request.form["investment_name11"]
    investment_name12 = request.form["investment_name12"]
    investment_name13 = request.form["investment_name13"]
    investment_name14 = request.form["investment_name14"]

    initial_investments = {
        investment_name0: investment0,
        investment_name1: investment1,
        investment_name2: investment2,
        investment_name3: investment3,
        investment_name4: investment4,
        investment_name5: investment5,
        investment_name6: investment6,
        investment_name7: investment7,
        investment_name8: investment8,
        investment_name9: investment9,
        investment_name10: investment10,
        investment_name11: investment11,
        investment_name12: investment12,
        investment_name13: investment13,
        investment_name14: investment14,
    }

    username = session['user']
    test = db.get_one(db.user_fm, "username", username)
    if not test:
        db.user_fm.insert_one({"username": username})

    
    db.user_fm.update_one({"username": username}, {"$set": {"initial_investments": initial_investments}}) 
    initial_investments_dict = db.get_one(db.user_fm, "username", username)['initial_investments']
    sum = 0
    for value in initial_investments_dict.values():
        if not is_digit(value):
            sum += int(0)
        else:
            sum += int(value)
    db.user_fm.update_one({"username": username}, {"$set": {"initial_investments_sum": sum}})
    return render_template('fm_constructor.html', form = fm_form)






@authorized.route('/salaries_data_saver', methods=['POST'])
@login_required
def salaries_data_saver():
    fm_form = fmForm(request.form)
    employee_position0 = request.form["employee_position0"]
    employee_position1 = request.form["employee_position1"]
    employee_position2 = request.form["employee_position2"]
    employee_position3 = request.form["employee_position3"]
    employee_position4 = request.form["employee_position4"]
    employee_position5 = request.form["employee_position5"]
    employee_position6 = request.form["employee_position6"]
    employee_position7 = request.form["employee_position7"]
    employee_position8 = request.form["employee_position8"]
    employee_position9 = request.form["employee_position9"]
    employee_position10 = request.form["employee_position10"]
    employee_position11 = request.form["employee_position11"]
    employee_position12 = request.form["employee_position12"]
    employee_position13 = request.form["employee_position13"]
    employee_position14 = request.form["employee_position14"]

    employee_salary0 = request.form["employee_salary0"]
    employee_salary1 = request.form["employee_salary1"]
    employee_salary2 = request.form["employee_salary2"]
    employee_salary3 = request.form["employee_salary3"]
    employee_salary4 = request.form["employee_salary4"]
    employee_salary5 = request.form["employee_salary5"]
    employee_salary6 = request.form["employee_salary6"]
    employee_salary7 = request.form["employee_salary7"]
    employee_salary8 = request.form["employee_salary8"]
    employee_salary9 = request.form["employee_salary9"]
    employee_salary10 = request.form["employee_salary10"]
    employee_salary11 = request.form["employee_salary11"]
    employee_salary12 = request.form["employee_salary12"]
    employee_salary13 = request.form["employee_salary13"]
    employee_salary14 = request.form["employee_salary14"]

    salaries = {
        employee_position0: employee_salary0,
        employee_position1: employee_salary1,
        employee_position2: employee_salary2,
        employee_position3: employee_salary3,
        employee_position4: employee_salary4,
        employee_position5: employee_salary5,
        employee_position6: employee_salary6,
        employee_position7: employee_salary7,
        employee_position8: employee_salary8,
        employee_position9: employee_salary9,
        employee_position10: employee_salary10,
        employee_position11: employee_salary11,
        employee_position12: employee_salary12,
        employee_position13: employee_salary13,
        employee_position14: employee_salary14,
    }

    username = session['user']
    test = db.get_one(db.user_fm, "username", username)
    if not test:
        db.user_fm.insert_one({"username": username})

    
    db.user_fm.update_one({"username": username}, {"$set": {"salaries": salaries}}) 
    salaries_dict = db.get_one(db.user_fm, "username", username)['salaries']
    sum = 0
    for value in salaries_dict.values():
        if not is_digit(value):
            sum += int(0)
        else:
            sum += int(value)
    db.user_fm.update_one({"username": username}, {"$set": {"salaries_sum": sum}})
    return render_template('fm_constructor.html', form = fm_form)








@authorized.route('/permanent_expenditures_data_saver', methods=['POST'])
@login_required
def permanent_expenditures_data_saver():
    fm_form = fmForm(request.form)
    permanent_expenditure_name0 = request.form["permanent_expenditure_name0"]
    permanent_expenditure_name1 = request.form["permanent_expenditure_name1"]
    permanent_expenditure_name2 = request.form["permanent_expenditure_name2"]
    permanent_expenditure_name3 = request.form["permanent_expenditure_name3"]
    permanent_expenditure_name4 = request.form["permanent_expenditure_name4"]
    permanent_expenditure_name5 = request.form["permanent_expenditure_name5"]
    permanent_expenditure_name6 = request.form["permanent_expenditure_name6"]
    permanent_expenditure_name7 = request.form["permanent_expenditure_name7"]
    permanent_expenditure_name8 = request.form["permanent_expenditure_name8"]
    permanent_expenditure_name9 = request.form["permanent_expenditure_name9"]
    permanent_expenditure_name10 = request.form["permanent_expenditure_name10"]
    permanent_expenditure_name11 = request.form["permanent_expenditure_name11"]
    permanent_expenditure_name12 = request.form["permanent_expenditure_name12"]
    permanent_expenditure_name13 = request.form["permanent_expenditure_name13"]
    permanent_expenditure_name14 = request.form["permanent_expenditure_name14"]

    permanent_expenditure0 = request.form["permanent_expenditure0"]
    permanent_expenditure1 = request.form["permanent_expenditure1"]
    permanent_expenditure2 = request.form["permanent_expenditure2"]
    permanent_expenditure3 = request.form["permanent_expenditure3"]
    permanent_expenditure4 = request.form["permanent_expenditure4"]
    permanent_expenditure5 = request.form["permanent_expenditure5"]
    permanent_expenditure6 = request.form["permanent_expenditure6"]
    permanent_expenditure7 = request.form["permanent_expenditure7"]
    permanent_expenditure8 = request.form["permanent_expenditure8"]
    permanent_expenditure9 = request.form["permanent_expenditure9"]
    permanent_expenditure10 = request.form["permanent_expenditure10"]
    permanent_expenditure11 = request.form["permanent_expenditure11"]
    permanent_expenditure12 = request.form["permanent_expenditure12"]
    permanent_expenditure13 = request.form["permanent_expenditure13"]
    permanent_expenditure14 = request.form["permanent_expenditure14"]

    permanent_expenditures = {
        permanent_expenditure_name0: permanent_expenditure0,
        permanent_expenditure_name1: permanent_expenditure1,
        permanent_expenditure_name2: permanent_expenditure2,
        permanent_expenditure_name3: permanent_expenditure3,
        permanent_expenditure_name4: permanent_expenditure4,
        permanent_expenditure_name5: permanent_expenditure5,
        permanent_expenditure_name6: permanent_expenditure6,
        permanent_expenditure_name7: permanent_expenditure7,
        permanent_expenditure_name8: permanent_expenditure8,
        permanent_expenditure_name9: permanent_expenditure9,
        permanent_expenditure_name10: permanent_expenditure10,
        permanent_expenditure_name11: permanent_expenditure11,
        permanent_expenditure_name12: permanent_expenditure12,
        permanent_expenditure_name13: permanent_expenditure13,
        permanent_expenditure_name14: permanent_expenditure14,
    }

    username = session['user']
    test = db.get_one(db.user_fm, "username", username)
    if not test:
        db.user_fm.insert_one({"username": username})

    
    db.user_fm.update_one({"username": username}, {"$set": {"permanent_expenditures": permanent_expenditures}}) 
    permanent_expenditures_dict = db.get_one(db.user_fm, "username", username)['permanent_expenditures']
    sum = 0
    for value in permanent_expenditures_dict.values():
        if not is_digit(value):
            sum += int(0)
        else:
            sum += int(value)
    db.user_fm.update_one({"username": username}, {"$set": {"permanent_expenditures_sum": sum}})
    return render_template('fm_constructor.html', form = fm_form)










@authorized.route('/variable_expenditures_data_saver', methods=['POST'])
@login_required
def variable_expenditures_data_saver():
    fm_form = fmForm(request.form)
    product1_name = request.form["product1_name"]
    product2_name = request.form["product2_name"]

    variable_expenditure_name0 = request.form["variable_expenditure_name0"]
    variable_expenditure_name1 = request.form["variable_expenditure_name1"]
    variable_expenditure_name2 = request.form["variable_expenditure_name2"]
    variable_expenditure_name3 = request.form["variable_expenditure_name3"]
    variable_expenditure_name4 = request.form["variable_expenditure_name4"]
    variable_expenditure_name5 = request.form["variable_expenditure_name5"]
    variable_expenditure_name6 = request.form["variable_expenditure_name6"]
    variable_expenditure_name7 = request.form["variable_expenditure_name7"]
    variable_expenditure_name8 = request.form["variable_expenditure_name8"]
    variable_expenditure_name9 = request.form["variable_expenditure_name9"]
    variable_expenditure_name10 = request.form["variable_expenditure_name10"]
    variable_expenditure_name11 = request.form["variable_expenditure_name11"]
    variable_expenditure_name12 = request.form["variable_expenditure_name12"]
    variable_expenditure_name13 = request.form["variable_expenditure_name13"]
    variable_expenditure_name14 = request.form["variable_expenditure_name14"]
    variable_expenditure_name15 = request.form["variable_expenditure_name15"]
    variable_expenditure_name16 = request.form["variable_expenditure_name16"]
    variable_expenditure_name17 = request.form["variable_expenditure_name17"]
    variable_expenditure_name18 = request.form["variable_expenditure_name18"]
    variable_expenditure_name19 = request.form["variable_expenditure_name19"]
    variable_expenditure_name20 = request.form["variable_expenditure_name20"]
    variable_expenditure_name21 = request.form["variable_expenditure_name21"]
    variable_expenditure_name22 = request.form["variable_expenditure_name22"]
    variable_expenditure_name23 = request.form["variable_expenditure_name23"]
    variable_expenditure_name24 = request.form["variable_expenditure_name24"]
    variable_expenditure_name25 = request.form["variable_expenditure_name25"]
    variable_expenditure_name26 = request.form["variable_expenditure_name26"]
    variable_expenditure_name27 = request.form["variable_expenditure_name27"]
    variable_expenditure_name28 = request.form["variable_expenditure_name28"]
    variable_expenditure_name29 = request.form["variable_expenditure_name29"]

    variable_expenditure0 = request.form["variable_expenditure0"]
    variable_expenditure1 = request.form["variable_expenditure1"]
    variable_expenditure2 = request.form["variable_expenditure2"]
    variable_expenditure3 = request.form["variable_expenditure3"]
    variable_expenditure4 = request.form["variable_expenditure4"]
    variable_expenditure5 = request.form["variable_expenditure5"]
    variable_expenditure6 = request.form["variable_expenditure6"]
    variable_expenditure7 = request.form["variable_expenditure7"]
    variable_expenditure8 = request.form["variable_expenditure8"]
    variable_expenditure9 = request.form["variable_expenditure9"]
    variable_expenditure10 = request.form["variable_expenditure10"]
    variable_expenditure11 = request.form["variable_expenditure11"]
    variable_expenditure12 = request.form["variable_expenditure12"]
    variable_expenditure13 = request.form["variable_expenditure13"]
    variable_expenditure14 = request.form["variable_expenditure14"]
    variable_expenditure15 = request.form["variable_expenditure15"]
    variable_expenditure16 = request.form["variable_expenditure16"]
    variable_expenditure17 = request.form["variable_expenditure17"]
    variable_expenditure18 = request.form["variable_expenditure18"]
    variable_expenditure19 = request.form["variable_expenditure19"]
    variable_expenditure20 = request.form["variable_expenditure20"]
    variable_expenditure21 = request.form["variable_expenditure21"]
    variable_expenditure22 = request.form["variable_expenditure22"]
    variable_expenditure23 = request.form["variable_expenditure23"]
    variable_expenditure24 = request.form["variable_expenditure24"]
    variable_expenditure25 = request.form["variable_expenditure25"]
    variable_expenditure26 = request.form["variable_expenditure26"]
    variable_expenditure27 = request.form["variable_expenditure27"]
    variable_expenditure28 = request.form["variable_expenditure28"]
    variable_expenditure29 = request.form["variable_expenditure29"]

    variable_expenditure1 = {
        variable_expenditure_name0: variable_expenditure0,
        variable_expenditure_name1: variable_expenditure1,
        variable_expenditure_name2: variable_expenditure2,
        variable_expenditure_name3: variable_expenditure3,
        variable_expenditure_name4: variable_expenditure4,
        variable_expenditure_name5: variable_expenditure5,
        variable_expenditure_name6: variable_expenditure6,
        variable_expenditure_name7: variable_expenditure7,
        variable_expenditure_name8: variable_expenditure8,
        variable_expenditure_name9: variable_expenditure9,
        variable_expenditure_name10: variable_expenditure10,
        variable_expenditure_name11: variable_expenditure11,
        variable_expenditure_name12: variable_expenditure12,
        variable_expenditure_name13: variable_expenditure13,
        variable_expenditure_name14: variable_expenditure14,
    }

    variable_expenditure2 = {
        variable_expenditure_name15: variable_expenditure15,
        variable_expenditure_name16: variable_expenditure16,
        variable_expenditure_name17: variable_expenditure17,
        variable_expenditure_name18: variable_expenditure18,
        variable_expenditure_name19: variable_expenditure19,
        variable_expenditure_name20: variable_expenditure20,
        variable_expenditure_name21: variable_expenditure21,
        variable_expenditure_name22: variable_expenditure22,
        variable_expenditure_name23: variable_expenditure23,
        variable_expenditure_name24: variable_expenditure24,
        variable_expenditure_name25: variable_expenditure25,
        variable_expenditure_name26: variable_expenditure26,
        variable_expenditure_name27: variable_expenditure27,
        variable_expenditure_name28: variable_expenditure28,
        variable_expenditure_name29: variable_expenditure29,
    }

    username = session['user']
    test = db.get_one(db.user_fm, "username", username)
    if not test:
        db.user_fm.insert_one({"username": username})

    
    db.user_fm.update_one({"username": username}, {"$set": {"variable_expenditure1": variable_expenditure1}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"variable_expenditure2": variable_expenditure2}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product1_name": product1_name}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product2_name": product2_name}}) 
    variable_expenditure1_dict = db.get_one(db.user_fm, "username", username)['variable_expenditure1']
    variable_expenditure2_dict = db.get_one(db.user_fm, "username", username)['variable_expenditure2']
    sum1 = 0
    for value in variable_expenditure1_dict.values():
        if not is_digit(value):
            sum1 += int(0)
        else:
            sum1 += int(value)
    db.user_fm.update_one({"username": username}, {"$set": {"variable_expenditure1_sum": sum1}})

    sum2 = 0
    for value in variable_expenditure2_dict.values():
        if not is_digit(value):
            sum2 += int(0)
        else:
            sum2 += int(value)
    db.user_fm.update_one({"username": username}, {"$set": {"variable_expenditure2_sum": sum1}})
    return render_template('fm_constructor.html', form = fm_form)







@authorized.route('/sales_data_saver', methods=['POST'])
@login_required
def sales_data_saver():
    fm_form = fmForm(request.form)
    
    product1_price = request.form["product1_price"]
    product1_expenditure_increase = request.form["product1_expenditure_increase"]
    product1_sales = request.form["product1_sales"]
    product1_sales_increase = request.form["product1_sales_increase"]
    
    product2_price = request.form["product2_price"]
    product2_expenditure_increase = request.form["product2_expenditure_increase"]
    product2_sales = request.form["product2_sales"]
    product2_sales_increase = request.form["product2_sales_increase"]
    
    
    username = session['user']
    test = db.get_one(db.user_fm, "username", username)
    if not test:
        db.user_fm.insert_one({"username": username})

    
    db.user_fm.update_one({"username": username}, {"$set": {"product1_price": product1_price}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product1_expenditure_increase": product1_expenditure_increase}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product1_sales": product1_sales}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product1_sales_increase": product1_sales_increase}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product2_price": product2_price}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product2_expenditure_increase": product2_expenditure_increase}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product2_sales": product2_sales}}) 
    db.user_fm.update_one({"username": username}, {"$set": {"product2_sales_increase": product2_sales_increase}}) 
    
    return render_template('fm_constructor.html', form = fm_form)







@authorized.route('/fm_constructor', methods=['GET', 'POST'])
@login_required
def fm_constructor():
    if request.method == "POST":
        username = session['user']
        mes1 = "У вас открылась главная страница, потому что пока нет доступа к той странице, на которую вы хотели зайти."
        mes2 = "Доступ к конструкторам будет поочередно открываться каждый день, для более качественного обучения"
        user = db.get_one(db.users, "username", username)
        user_status = user['payment_status']
        if user_status == "course-pro" or user_status == "course-standard" or user_status == "beginner":
            return render_template(f"index.html", message1 = mes1, message2 = mes2)
    fm_form =  fmForm(request.form)
    if request.method == 'GET':
        logging.error("this war 1")
        # return render_template(f"index.html", message1 = mes1, message2 = mes2)
        return render_template('fm_constructor.html', form = fm_form)


# @authorized.route('/<template>')
# @login_required
# def route_template(template):




#     try:

#         if not template.endswith( '.html' ):
#             template += '.html'

#         # Detect the current page
#         segment = get_segment( request )

#         # Serve the file (if exists) from app/templates/FILE.html
#         return render_template( template, segment=segment )

#     except TemplateNotFound:
#         return render_template('page-404.html'), 404
    
#     except:
#         return render_template('page-500.html'), 500

# Helper - Extract current page name from request 
def get_segment( request ): 

    try:

        segment = request.path.split('/')[-1]

        if segment == '':
            segment = 'index'

        return segment    

    except:
        return None  

# ## Login & Registration

# @blueprint.route('/login', methods=['GET', 'POST'])
# def login():
#     login_form = LoginForm(request.form)
#     if 'login' in request.form:
        
#         # read form data
#         username = request.form['username']
#         password = request.form['password']

#         # Locate user
#         user = User.query.filter_by(username=username).first()
        
#         # Check the password
#         if user and verify_pass( password, user.password):

#             login_user(user)
#             return redirect(url_for('base_blueprint.route_default'))

#         # Something (user or pass) is not ok
#         return render_template( 'accounts/login.html', msg='Wrong user or password', form=login_form)

#     if not current_user.is_authenticated:
#         return render_template( 'accounts/login.html',
#                                 form=login_form)
#     # return redirect(url_for('base_blueprint.index'))
#     return render_template('index.html')

# @blueprint.route('/register', methods=['GET', 'POST'])
# def register():
#     login_form = LoginForm(request.form)
#     create_account_form = CreateAccountForm(request.form)
#     if 'register' in request.form:

#         username  = request.form['username']
#         email     = request.form['email'   ]

#         # Check usename exists
#         user = User.query.filter_by(username=username).first()
#         if user:
#             return render_template( 'accounts/register.html', 
#                                     msg='Username already registered',
#                                     success=False,
#                                     form=create_account_form)

#         # Check email exists
#         user = User.query.filter_by(email=email).first()
#         if user:
#             return render_template( 'accounts/register.html', 
#                                     msg='Email already registered', 
#                                     success=False,
#                                     form=create_account_form)

#         # else we can create the user
#         user = User(**request.form)
#         db.session.add(user)
#         db.session.commit()

#         return render_template( 'accounts/register.html', 
#                                 msg='User created please <a href="/login">login</a>', 
#                                 success=True,
#                                 form=create_account_form)

#     else:
#         return render_template( 'accounts/register.html', form=create_account_form)

# @blueprint.route('/logout')
# def logout():
#     logout_user()
#     return redirect(url_for('base_blueprint.login'))

## Errors

# @login_manager.unauthorized_handler
# def unauthorized_handler():
#     return render_template('page-403.html'), 403

@authorized.errorhandler(403)
def access_forbidden(error):
    return render_template('page-403.html'), 403

@authorized.errorhandler(404)
def not_found_error(error):
    return render_template('page-404.html'), 404

@authorized.errorhandler(500)
def internal_error(error):
    return render_template('page-500.html'), 500





def replace_string(filename, oldtext, newtext, newfilename):
    doc = Document(filename)
    for p in doc.paragraphs:
        if oldtext in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if oldtext in inline[i].text:
                    text = inline[i].text.replace(oldtext, newtext)
                    inline[i].text = text
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if oldtext in cell.text:
                    # inline = cell.runs
                    # # Loop added to work with runs (strings with same style)
                    # for i in range(len(inline)):
                    #     if oldtext in inline[i].text:
                    text = cell.text.replace(oldtext, newtext)
                    cell.text = text

    doc.save(newfilename)
    return 1